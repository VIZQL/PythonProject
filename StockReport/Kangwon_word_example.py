# 가장 기본적인 기능(문서열기, 저장, 글자쓰기 등등)
from docx import Document

# 문단 정렬 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 테이블 정렬 
from docx.enum.table import WD_TABLE_ALIGNMENT


# 문자 스타일 변경 
from docx.enum.style import WD_STYLE_TYPE

# 이미지 삽입
# Cm 과 Inch 단위를 사용하기 위한 모듈 
from docx.shared import Cm,Inches

# 한글 폰트 선 문단 입력, 후 폰트적용을 위한 import 
from docx.oxml.ns import qn

# 문자 크기를 변경하기 위해 폰트 크기와 관련된 class 를 import 
from docx.shared import Pt

# RGB color 를 사용하는 메쏘드
from docx.shared import RGBColor

# word 문서를 PDF 로 바꾸기 위한 모듈 
import aspose.words as aw
# pip install aspose-words

# 목차를 만들기 위한 import 
import win32com.client
import os



# 새 워드 문서 만들기 
doc = Document()

# 현재 작업경로에 저장 
# doc.save('new_try.docx')

# 저장된 문서 불러오기 
# doc = Document('절대경로\docx') # 절대 경로 불러오기 
# doc = Document('C:\PYTHONWORKSPACE\\new_try.docx')

# # 1. 제목 넣기 
# doc.add_heading('가장 큰 제목 (아래에 밑줄)', level = 0)
# doc.add_heading('제목 크기 H1', level = 1)
# doc.add_heading('제목 크기 H2', level = 2)
# doc.add_heading('제목 크기 H3', level = 3)
# doc.add_heading('제목 크기 H4', level = 4)

# # 2. 문단 추가하기 
# doc.add_paragraph('여기에 원하는 텍스트를 마음껏 입력하시면 됩니다')

# # 3. 문단에 문자 추가하기 
# p = doc.add_paragraph('두번째 문단 : 여기에 원하는 텍스트를 마음껏 입력하시면 됩니다')

# # 굵은 글씨(Bold) 적용
# p.add_run('문단에 굵은 글자 추가').bold = True

# # 기울임꼴 (italic) 적용
# p.add_run('문단에 굵은 글자 추가').italic = True

# # 밑줄  (Underline) 적용
# p.add_run('문단에 밑줄 추가').underline = True


# # 이미지 삽입 

# # 사진의 크기를 Cm 단위로 설정하여 삽입 
# doc.add_picture('0HR_BEMF파형.JPG', width = Cm(16), height = Cm(9))

# # 사진의 크기를 Inch 단위로 설정하여 삽입 
# # doc.add_picture('0HR_BEMF파형.JPG', width = Inches(4), height = Inches(9))


# # 표 삽입 - 2행 3열의 표 만들기 
# table = doc.add_table(rows = 2, cols = 3)

# # 만든 표의 스타일을 가장 기본 스타일인 'Table Grid' 로 설정 
# table.style = doc.styles['Table Grid']

# # 표의 첫 행을 리스트로 가져오기 
# first_row = table.rows[0].cells

# # 첫 행의 각 열들에 접근해서 값 입력 
# first_row[0].text = 'a'
# first_row[1].text = 'b'
# first_row[2].text = 'c'

# # 표의 두번째 행을 리스트로 가져온 후 , 각 셀에 값 입력 
# second_row = table.rows[1].cells
# second_row[0].text = 'd'
# second_row[1].text = 'e'
# second_row[2].text = 'f'

# # 행 추가하기  
# row_cells = table.add_row().cells

# # 열 추가하기 
# col_cells = table.add_column(width=Cm(2)).cells


# 현재 작업경로에 저장 
# doc.save('new_try.docx')

# # 파일 읽어와서 작업하기 

# doc = Document('예제 문서.docx')

# for i, paragraph in enumerate(doc.paragraphs):
#     print(str(i+1)+ ": " + paragraph.text)

# # 추가하고 싶은 문단 ( *실제 인덱스는 0부터 시작하므로 문단번호보다 +1 해주어야함)
# p = doc.paragraphs[4]

# p.add_run('문단에 굵은 글자 추가')

# # 3번째 문장 다음에 삽입 (*paragraphs[3] 은 4번째 문장을 의미한다 )
# doc.paragraphs[3].insert_paragraph_before("문장을 삽입한다.")

# # 문서 안의 모든 표를 가져옴 
# tables = doc.tables

# # 가장 처음 표의 첫행, 첫 열의 첫 문단 내용 가져오기 
# print(tables[0].rows[0].cells[0].paragraphs[0].text)

# # 표의 모든값에 접근 
# table = doc.tables[0]

# for row in table.rows:
#     for cell in row.cells:
#         for para in cell.paragraphs:
#             print(para.text)


# # 글자 찾아서 수정하기 

# for row in table.rows:
#     for cell in row.cells:
#         for para in cell.paragraphs:
#             if(para.text == "하나"):
#                 para.add_run('<-- 찾았다 하나')
#             # print(para.text)

# # 확인을 위해서 문서 저장
# doc.save("예제 문서.docx")

# # 새로운 문서 만들기 
# doc = Document()

# # 스타일 적용하기 
# style = doc.styles['Normal']
# font = style.font
# font.name = 'Arial'

# para = doc.add_paragraph('Some Text\n')

# # 선문단 입력후 스타일 적용
# para.add_run('코딩유치원에 오신 것을 환영합니다.').bold = True

# run = doc.paragraphs[0].runs[0]

# run.font.name = 'Arial'

# doc = Document()

# style = doc.styles['Normal']
# style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은고딕')
# style.font.name = '맑은고딕'
# style.font.size = Pt(8)

# para = doc.add_paragraph('맑은 고딕체 테스트')

# doc.save('text.docx')

# doc = Document()

# para = doc.add_paragraph('이 글자의 크기를 바꿔봅시다')

# # para = doc.add_paragraph('두번째 문장입니다.')

# # p = para.add_run('\n')
# # p = para.add_run('두번째 문장의 run 입니다')

# # 첫번째 문단의 문장(run) 들을 리스트로 받기 
# para1 = doc.paragraphs[0].runs

# para2 = doc.add_paragraph('글자 색깔을 바꿔봅시다')
# run = para2.runs[0]
# font = run.font

# # RGB 컬러를 각각 16진수로 표현 (R,G,B)
# font.color.rgb = RGBColor(0xFF,0x24,0xE9)



# # # for 문을 이용해서 
# # for run in para1:
# #     run.font.size = Pt(20)

# doc.save('text1.docx')

# doc = Document('예제 문서.docx')

# # 왼쪽정렬 
# paragraph1 = doc.paragraphs[1]
# paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT

# # 가운데 정렬
# paragraph2 = doc.paragraphs[2]
# paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# # 오른쪽 정렬
# paragraph3 = doc.paragraphs[3]
# paragraph3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# # 양쪽 정렬
# paragraph4 = doc.paragraphs[4]
# paragraph4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# # 텍스트 배분 (글자를 흩어서 배치)
# paragraph_last = doc.paragraphs[-1] # 마지막 문단
# paragraph_last.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
# paragraph_last.alignment = WD_PARAGRAPH_ALIGNMENT.DITRIBUTE # 컴파일 에러남

# 현재 작업경로에 저장 
# doc.save('정렬 예제.docx')

# 테이블 셀 정렬 
# # LEFT : 왼쪽정렬,  CENTER :  가운데 정렬, RIGHT : 오른쪽정렬 
# doc.tables[0].rows[0].cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT
# doc.tables[0].rows[0].cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
# doc.tables[0].rows[0].cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.RIGHT


# # doc.save('TryToc1.docx')

# 목차 생성 
word = win32com.client.DispatchEx("Word.Application")
word.Visible = False
doc = word.Documents.Open('C:/PYTHONWORKSPACE/new_try.docx', Encoding='gbk')
doc.TablesofContents(1).Update()
doc.Close(SaveChange=True)
word.Quit()


# pdf 파일로 변환
# doc = aw.Document('정렬 예제.docx')

# doc.save('정렬예제.pdf')



# doc.save('정렬 예제.pdf')


