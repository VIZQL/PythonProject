# 가장 기본적인 기능(문서열기, 저장, 글자쓰기 등등)
from docx import Document

# 문단 정렬 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 테이블 정렬 
from docx.enum.table import WD_TABLE_ALIGNMENT

# 테이블 셀 색 넣기
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# 문자 스타일 변경 
from docx.enum.style import WD_STYLE_TYPE

# 이미지 삽입
# Cm 과 Inch 단위를 사용하기 위한 모듈 
from docx.shared import Cm,Inches

# 한글 폰트 선 문단 입력, 후 폰트적용을 위한 import 
from docx.oxml.ns import qn

# 문자 크기를 변경하기 위해 폰트 크기와 관련된 class 를 import 
from docx.shared import Pt

# RGB color 를 사용하는 메쏘드
from docx.shared import RGBColor

# word 문서를 PDF 로 바꾸기 위한 모듈 
# import aspose.words as aw
# pip install aspose-words

# 목차를 만들기 위한 import 
import win32com.client
import os
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 현재 날짜 가져오기
# import datetime
# from datetime import datetime

# 내 모듈 import 하기
from Kangwon_News import*
from Kangwon_stocks_info import*
from stock_report import*
from Kangwon_schedule import*
from Kangwon_raw import*
from Kiwoom_condition import*

import Kangwon_stocks_info
import stock_report
import Kangwon_News
import Kangwon_schedule
import Kangwon_raw
import Kiwoom_condition

# pdf 로 변환하기 위한 메소드
from docx2pdf import convert


# -*- coding: utf-8 -*

# 하이퍼 링크를 위한 함수 
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


# 셀 컬러 정의하기 

def cell_color(firstrow):
    for i in range (len(firstrow)):
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="D6E8F6"/>'.format(nsdecls('w')))
        firstrow[i]._tc.get_or_add_tcPr().append(shading_elm_1)

# 원하는 종목 특징주에서 뽑아오고 특징주 list 에서 삭제 

def stock_news_return(name, news_list, link_list):
    name_re = re.compile(name)
    news_result = []
    link_result = []

    for index, news in enumerate(news_list):
        m = name_re.search(news) # 주어진 문자열 중에 일치하는게 있는지 확인  
        
        if m:
            news_result.append(news_list[index])
            link_result.append(link_list[index])
            news_list.pop(index)
            link_list.pop(index)
        else:
            continue

    if(news_result): # 매칭 종목이 있을때 
        return news_result, link_result
    else: # 매칭 종목이 없을 때 
        return [""], [""]


################################# 여기서부터 문서작업 #######################################

# 새 워드 문서 만들기 
## test
with open("테스트.txt", "w" ) as f:
    f.write("\n")
    f.write("-"*100 + "\n")
    f.write(str(datetime.datetime.now().strftime('%Y-%m-%d %H:%M')) + " :테스트 ")


# doc = Document()
doc = Document('C:\PYTHONWORKSPACE\\TP7_YR.docx')

# today_time = datetime.today().strftime("%Y%m%d")   
today_time = datetime.date.today().strftime("%Y%m%d")  

# 헤딩 폰트 사이즈 변경
font = doc.styles['Title'].font
font.name = 'Times New Roman'
font.size = Pt(20)

font = doc.styles['Heading 1'].font
font.name = 'Times New Roman'
font.color.rgb = RGBColor(0,0,0)


# 글자 사이즈 변경 
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(8)

# 1. 제목 넣기 
doc.add_heading('Y&R report_{} '.format(today_time), level = 0)

# 2. 문단 추가하기 
p = doc.add_paragraph()
p.add_run('Y&R 리포트는 저작권의 보호를 받습니다. 작성자 허가없이 다른 사람에게 공유는 엄격히 금지됩니다. 지정된 수신자 외에 다른 사람에게 전달되는 것이 적발될 경우 민형사상의 책임을 질 수 있습니다.').underline = True
p.runs[0].font.size = Pt(8)

################################ 목차 생성 ########################################
paragraph = doc.add_paragraph()
run = paragraph.add_run()

fldChar = OxmlElement('w:fldChar')  # creates a new element
fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:t')
fldChar3.text = "Right-click to update field."
fldChar2.append(fldChar3)

fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')

r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar4)
p_element = paragraph._p

################################ 다음페이지로 넘어가기 ########################################
doc.add_page_break()

################################ 1. Daily Comment ########################################
doc.add_heading('1. Daily Comment : ', level = 1)


################################ 2. 주요 국가 지수 ########################################
doc.add_heading('2. 주요 국가 지수 : ', level = 1)

# stocks_info 수행 

indice_name, indice_number, indice_percent =  Kangwon_stocks_info.scrape_major_indice()

# 표 삽입 - 6행 3열의 표 만들기 
table = doc.add_table(rows = 6, cols = 3)

# 만든 표의 스타일을 가장 기본 스타일인 'Table Grid' 로 설정 
table.style = doc.styles['Table Grid']

# 폰트 사이즈 재 변경 
style.font.size = Pt(10)

# 표의 첫 행을 리스트로 가져오기 
first_row = table.rows[0].cells

cell_color(first_row)

first_row[0].text = ''
first_row[1].text = '지수'
first_row[2].text = '증감율(%)'



for i in range(0,5):
    data = table.rows[i+1].cells
    data[0].text = indice_name[i]
    data[1].text = indice_number[i]
    data[2].text = indice_percent[i]
    if (float(indice_percent[i].strip("%")) > 0):
        data[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0x00,0x00)
    elif (float(indice_percent[i].strip("%")) < 0):
        data[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x00,0xFF)
    else:
        pass


################################ 3. 오전장 특징주  ########################################
doc.add_heading('3. 오전장 특징주 : ', level = 1)


try:
    title1, stext1, text1 = Kangwon_stocks_info.notable_stocks("오전장 특징주★(코스닥)")
    title2, stext2, text2 = Kangwon_stocks_info.notable_stocks("오전장 특징주★(코스피)")
except:pass


p = doc.add_paragraph()
p.add_run("<코스닥>").bold = True
p.runs[0].font.size = Pt(10)

# 표 삽입 - 1행 2열의 표 만들기 
table = doc.add_table(rows = 1, cols = 2)

# 표 사이즈 조정
table.columns[0].width = Cm(1)
table.columns[1].width = Cm(6)


# 만든 표의 스타일을 가장 기본 스타일인 'Table Grid' 로 설정 
table.style = doc.styles['Table Grid']


first_row = table.rows[0].cells

cell_color(first_row)

# 표 사이즈 재 조정
first_row[0].width = Cm(2)
first_row[1].width = Cm(12)

first_row[0].text = '종목'
first_row[1].text = '내용'

pre_temp = 0

for i in range(len(title1)*2):
    data = table.add_row().cells

    if i%2 == 0:
        data[0].text = title1[i//2]
        data[1].text = stext1[i//2]
    else:
        data[1].text = text1[i//2]
        data[0].merge(pre_temp)

    pre_temp = data[0]

print("3. 코스닥 완료")


p = doc.add_paragraph()
p.add_run("\n").bold = True
p.add_run("<코스피>").bold = True
p.runs[0].font.size = Pt(10)

# 표 삽입 - 1행 2열의 표 만들기 
table = doc.add_table(rows = 1, cols = 2)

# 표 사이즈 조정
table.columns[0].width = Cm(1)
table.columns[1].width = Cm(6)


# 만든 표의 스타일을 가장 기본 스타일인 'Table Grid' 로 설정 
table.style = doc.styles['Table Grid']


first_row = table.rows[0].cells

cell_color(first_row)

# 표 사이즈 재 조정
first_row[0].width = Cm(2)
first_row[1].width = Cm(12)

first_row[0].text = '종목'
first_row[1].text = '내용'

pre_temp = 0

for i in range(len(title2)*2):
    data = table.add_row().cells

    if i%2 ==0:
        data[1].text = stext2[i//2]
        data[0].text = title2[i//2]
    else:
        data[1].text = text2[i//2]
        data[0].merge(pre_temp)

    pre_temp = data[0]

print("3. 코스피 완료")



################################ 4. 시장 주도 종목 정리 ########################################
doc.add_heading('4. 시장 주도 종목 정리 : ', level = 1)


try:
    Kiwoom_condition.getConditionKiwoom()
    title_list, link_list = scrape_stocks_info_saghan(Kiwoom_condition.sanghan_name)
except:pass


### 상한가 종목 정리 
p = doc.add_paragraph()
p.add_run("<상한가>").bold = True
p.runs[0].font.size = Pt(10)

# 표 삽입 - 1행 3열의 표 만들기 
table = doc.add_table(rows = 1, cols = 3)
table.columns[0].width = Cm(4)
table.columns[1].width = Cm(4)
table.columns[2].width = Cm(6)

# 만든 표의 스타일을 가장 기본 스타일인 'Table Grid' 로 설정 
table.style = doc.styles['Table Grid']

first_row = table.rows[0].cells

cell_color(first_row)

first_row[0].text = '종목'
first_row[1].text = '거래대금(10억)'
first_row[2].text = '비고'

first_row[0].width = Cm(4)
first_row[1].width = Cm(4)
first_row[2].width = Cm(6)

try:
    for index in range(len(Kiwoom_condition.sanghan_name)):

        data = table.add_row().cells
        tmp_name =  Kiwoom_condition.sanghan_name[index]
        data[0].text = tmp_name
        data[1].text = Kiwoom_condition.sanghan_tramount[index]
        data[2].text = title_list[index]
        p = data[2].paragraphs[0]
        p.add_run("\n")
        add_hyperlink(p, link_list[index], link_list[index])
except:pass

doc.add_paragraph('\n') # 한칸 띄우기 

### 거래대금 상위 종목 정리 

p = doc.add_paragraph()
p.add_run("<거래대금 상위 종목>").bold = True
p.runs[0].font.size = Pt(10)

# 표 삽입 - 1행 4열의 표 만들기 
table = doc.add_table(rows = 1, cols = 4)

table.columns[0].width = Cm(1)
table.columns[1].width = Cm(2)
table.columns[2].width = Cm(2)
table.columns[3].width = Cm(2)

# 만든 표의 스타일을 가장 기본 스타일인 'Table Grid' 로 설정 
table.style = doc.styles['Table Grid']

first_row = table.rows[0].cells

cell_color(first_row)

first_row[0].text = '순위'
first_row[1].text = '종목'
first_row[2].text = '등락률 (%)'
first_row[3].text = '거래대금 (10억)'

first_row[0].width = Cm(1)
first_row[1].width = Cm(3)
first_row[2].width = Cm(3)
first_row[3].width = Cm(3)

try:
    for index in range(len(Kiwoom_condition.trTop_name)):

        data = table.add_row().cells
        data[0].text = str(index+1)
        data[1].text = Kiwoom_condition.trTop_name[index]
        data[2].text = Kiwoom_condition.trTop_percent[index] + "%"
        if (float(Kiwoom_condition.trTop_percent[index]) > 0):
            data[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0x00,0x00)
        elif (float(Kiwoom_condition.trTop_percent[index]) < 0):
            data[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x00,0xFF)
        else:
            pass
        data[3].text = Kiwoom_condition.trTop_amount[index]
except:pass


### 주목 받은 종목 

title_list, link_list = scrape_stocks_info_saghan(Kiwoom_condition.attention_name)

doc.add_paragraph('\n') # 한칸 띄우기 

p = doc.add_paragraph()
p.add_run("<주목받은 종목>").bold = True
p.runs[0].font.size = Pt(10)

# 표 삽입 - 1행 4열의 표 만들기 
table = doc.add_table(rows = 1, cols = 4)

table.columns[0].width = Cm(3)
table.columns[1].width = Cm(3)
table.columns[2].width = Cm(3)
table.columns[3].width = Cm(4)

# 만든 표의 스타일을 가장 기본 스타일인 'Table Grid' 로 설정 
table.style = doc.styles['Table Grid']

first_row = table.rows[0].cells

cell_color(first_row)

first_row[0].text = '종목'
first_row[1].text = '등락률 (%)'
first_row[2].text = '거래대금 (10억)'
first_row[3].text = '비고'

first_row[0].width = Cm(3)
first_row[1].width = Cm(3)
first_row[2].width = Cm(3)
first_row[3].width = Cm(4)

try:
    for index in range(len(Kiwoom_condition.attention_name)):

        data = table.add_row().cells
        data[0].text = Kiwoom_condition.attention_name[index]
        data[1].text = Kiwoom_condition.attention_percent[index] + "%"
        if (float(Kiwoom_condition.attention_percent[index]) > 0):
            data[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0x00,0x00)
        elif (float(Kiwoom_condition.attention_percent[index]) < 0):
            data[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x00,0xFF)
        else:
            pass
        data[2].text = Kiwoom_condition.attention_amount[index]
        data[3].text = title_list[index]
        p = data[3].paragraphs[0]
        p.add_run("\n")
        add_hyperlink(p, link_list[index], link_list[index])
except:pass

################################ 5. 그 외 특징종목 정리 ########################################

doc.add_heading('5. 그 외 특징주 정리 : ', level = 1)

title_list, link_list = Kangwon_stocks_info.scrape_stocks_info()

for index in range(0,len(title_list)):
    p = doc.add_paragraph(title_list[index])
    p.add_run('\n')
    add_hyperlink(p, link_list[index], link_list[index])

print("5. 그 외 특징종목 정리 완료")


################################ 6. 오전 주요 뉴스  ####################################### 
doc.add_heading('6. 오전 주요 뉴스: ', level = 1)

Kangwon_News.scrape_headline_news()

title_list, link_list = Kangwon_News.scrape_headline_news()
title_list_eco, link_list_eco = Kangwon_News.scrape_headline_news_eco()
title_list_bell, link_list_bell = Kangwon_News.scrape_headline_news_thebell()
title_list_guru, link_list_guru = Kangwon_News.scrape_headline_news_guru()


style.font.size = Pt(10)


for i in range(0,len(title_list)):
    if(i == 0*Kangwon_News.limit_news):
        p = doc.add_paragraph(30*"-"+"<정치>"+30*"-")
    elif(i == 1*Kangwon_News.limit_news):
        p = doc.add_paragraph(30*"-"+"<경제>"+30*"-")
    elif(i == 2*Kangwon_News.limit_news):
        p = doc.add_paragraph(30*"-"+"<세계>"+30*"-")

    p = doc.add_paragraph()
    p.add_run(title_list[i])
    p.add_run('\n')
    add_hyperlink(p,link_list[i],link_list[i])

for i in range(0,len(title_list_eco)):
    if(i == 0*Kangwon_News.limit_news):
        p = doc.add_paragraph(30*"-"+"<금융>"+30*"-")
    elif(i == 1*Kangwon_News.limit_news):
        p = doc.add_paragraph(30*"-"+"<증권>"+30*"-")
    elif(i == 2*Kangwon_News.limit_news):
        p = doc.add_paragraph(30*"-"+"<산업>"+30*"-")

    p = doc.add_paragraph()
    p.add_run(title_list_eco[i])
    p.add_run('\n')
    add_hyperlink(p,link_list_eco[i],link_list_eco[i])

for i in range(0,len(title_list_bell)):
    if(i == 0*Kangwon_News.limit_news_bell):
        p = doc.add_paragraph(30*"-"+"<기업-1>"+30*"-")

    p = doc.add_paragraph()
    p.add_run(title_list_bell[i])
    p.add_run('\n')
    add_hyperlink(p,link_list_bell[i],link_list_bell[i])

for i in range(0,len(title_list_guru)):
    if(i == 0*Kangwon_News.limit_news_guru):
        p = doc.add_paragraph(30*"-"+"<기업-2>"+30*"-")

    p = doc.add_paragraph()
    p.add_run(title_list_guru[i])
    p.add_run('\n')
    add_hyperlink(p,link_list_guru[i],link_list_guru[i])

print("6. 주요 뉴스 완료")

################################ 7. 관심 차트  ####################################### 
doc.add_heading('7. 관심 차트: ', level = 1)


p = doc.add_paragraph()
p.add_run("<이평선 매매>").bold = True
p.add_run(": 최근 의미있는 상승 후 눌림(이평선 근접) 종목")
p.runs[0].font.size = Pt(10)


# 표 삽입 - 1행 2열의 표 만들기 
table = doc.add_table(rows = 1, cols = 2)

table.columns[0].width = Cm(1)
table.columns[1].width = Cm(1)

# 만든 표의 스타일을 가장 기본 스타일인 'Table Grid' 로 설정 
table.style = doc.styles['Table Grid']

first_row = table.rows[0].cells

cell_color(first_row)

# 표 사이즈 재 조정

first_row[0].text = '분류'
first_row[1].text = '종목'

# 10일 선 
data = table.add_row().cells
data[0].text = '10일선'
for name in Kiwoom_condition.chart8sun:
    p = data[1].paragraphs[0]
    if(name == Kiwoom_condition.chart8sun[-1]):
        p.add_run(name)
    else:
        p.add_run(name +', ')

# 20일선 
data = table.add_row().cells 
data[0].text = '20일선'
for name in Kiwoom_condition.chart20sun:
    p = data[1].paragraphs[0]
    if(name == Kiwoom_condition.chart20sun[-1]):
        p.add_run(name)
    else:
        p.add_run(name +', ')

# 45일선 
data = table.add_row().cells 
data[0].text = '45일선'
for name in Kiwoom_condition.chart45sun:
    p = data[1].paragraphs[0]
    if(name == Kiwoom_condition.chart45sun[-1]):
        p.add_run(name)
    else:
        p.add_run(name +', ')

print("7. 관심 차트 완료 ")


################################ 현재 작업경로에 저장  ####################################### 
# doc.save('C:\\PYTHONWORKSPACE\\webscraping_basic\\webscraping_project\\2022\\kikawo_SD_리포트_{}.docx'.format(today_time))
doc.save('C:\\PYTHONWORKSPACE\\webscraping_basic\\webscraping_project\\2022\\Y&R_리포트_{}_오전장.docx'.format(today_time))

time.sleep(3)
################################ 목차 업데이트 ####################################### 
word = win32com.client.DispatchEx("Word.Application")
doc = word.Documents.Open('C:\\PYTHONWORKSPACE\\webscraping_basic\\webscraping_project\\2022\\Y&R_리포트_{}_오전장.docx'.format(today_time))
doc.TablesOfContents(1).Update()
doc.Close(SaveChanges=True)
word.Quit()

time.sleep(5)
################################ pdf 변환 #################################################### 
inputFile = "C:\\PYTHONWORKSPACE\\webscraping_basic\\webscraping_project\\2022\\Y&R_리포트_{}_오전장.docx".format(today_time)
outputFile = "C:\\PYTHONWORKSPACE\\webscraping_basic\\webscraping_project\\2022\\Y&R_리포트_{}_오전장.pdf".format(today_time)
file = open(outputFile, "w")
file.close()

convert(inputFile, outputFile)


# 저장된 문서 불러오기 
# doc = Document('절대경로\docx') # 절대 경로 불러오기 
# doc = Document('C:\PYTHONWORKSPACE\\new_try.docx')


# doc.add_heading('제목 크기 H1', level = 1)
# doc.add_heading('제목 크기 H2', level = 2)
# doc.add_heading('제목 크기 H3', level = 3)
# doc.add_heading('제목 크기 H4', level = 4)

# # 2. 문단 추가하기 
# doc.add_paragraph('여기에 원하는 텍스트를 마음껏 입력하시면 됩니다')

# # 3. 문단에 문자 추가하기 
# p = doc.add_paragraph('두번째 문단 : 여기에 원하는 텍스트를 마음껏 입력하시면 됩니다')

# # 굵은 글씨(Bold) 적용
# p.add_run('문단에 굵은 글자 추가').bold = True

# # 기울임꼴 (italic) 적용
# p.add_run('문단에 굵은 글자 추가').italic = True

# # 밑줄  (Underline) 적용
# p.add_run('문단에 밑줄 추가').underline = True


# # 이미지 삽입 

# # 사진의 크기를 Cm 단위로 설정하여 삽입 
# doc.add_picture('0HR_BEMF파형.JPG', width = Cm(16), height = Cm(9))

# # 사진의 크기를 Inch 단위로 설정하여 삽입 
# # doc.add_picture('0HR_BEMF파형.JPG', width = Inches(4), height = Inches(9))


# # 표 삽입 - 2행 3열의 표 만들기 
# table = doc.add_table(rows = 2, cols = 3)

# # 만든 표의 스타일을 가장 기본 스타일인 'Table Grid' 로 설정 
# table.style = doc.styles['Table Grid']

# # 표의 첫 행을 리스트로 가져오기 
# first_row = table.rows[0].cells

# # 첫 행의 각 열들에 접근해서 값 입력 
# first_row[0].text = 'a'
# first_row[1].text = 'b'
# first_row[2].text = 'c'

# # 표의 두번째 행을 리스트로 가져온 후 , 각 셀에 값 입력 
# second_row = table.rows[1].cells
# second_row[0].text = 'd'
# second_row[1].text = 'e'
# second_row[2].text = 'f'

# # 행 추가하기  
# row_cells = table.add_row().cells

# # 열 추가하기 
# col_cells = table.add_column(width=Cm(2)).cells


# 현재 작업경로에 저장 
# doc.save('new_try.docx')

# # 파일 읽어와서 작업하기 

# doc = Document('예제 문서.docx')

# for i, paragraph in enumerate(doc.paragraphs):
#     print(str(i+1)+ ": " + paragraph.text)

# # 추가하고 싶은 문단 ( *실제 인덱스는 0부터 시작하므로 문단번호보다 +1 해주어야함)
# p = doc.paragraphs[4]

# p.add_run('문단에 굵은 글자 추가')

# # 3번째 문장 다음에 삽입 (*paragraphs[3] 은 4번째 문장을 의미한다 )
# doc.paragraphs[3].insert_paragraph_before("문장을 삽입한다.")

# # 문서 안의 모든 표를 가져옴 
# tables = doc.tables

# # 가장 처음 표의 첫행, 첫 열의 첫 문단 내용 가져오기 
# print(tables[0].rows[0].cells[0].paragraphs[0].text)

# # 표의 모든값에 접근 
# table = doc.tables[0]

# for row in table.rows:
#     for cell in row.cells:
#         for para in cell.paragraphs:
#             print(para.text)


# # 글자 찾아서 수정하기 

# for row in table.rows:
#     for cell in row.cells:
#         for para in cell.paragraphs:
#             if(para.text == "하나"):
#                 para.add_run('<-- 찾았다 하나')
#             # print(para.text)

# # 확인을 위해서 문서 저장
# doc.save("예제 문서.docx")

# # 새로운 문서 만들기 
# doc = Document()

# # 스타일 적용하기 
# style = doc.styles['Normal']
# font = style.font
# font.name = 'Arial'

# para = doc.add_paragraph('Some Text\n')

# # 선문단 입력후 스타일 적용
# para.add_run('코딩유치원에 오신 것을 환영합니다.').bold = True

# run = doc.paragraphs[0].runs[0]

# run.font.name = 'Arial'

# doc = Document()

# style = doc.styles['Normal']
# style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은고딕')
# style.font.name = '맑은고딕'
# style.font.size = Pt(8)

# para = doc.add_paragraph('맑은 고딕체 테스트')

# doc.save('text.docx')

# doc = Document()


# para = doc.add_paragraph('이 글자의 크기를 바꿔봅시다')

# # para = doc.add_paragraph('두번째 문장입니다.')

# # p = para.add_run('\n')
# # p = para.add_run('두번째 문장의 run 입니다')

# # 첫번째 문단의 문장(run) 들을 리스트로 받기 
# para1 = doc.paragraphs[0].runs

# para2 = doc.add_paragraph('글자 색깔을 바꿔봅시다')
# run = para2.runs[0]
# font = run.font

# # RGB 컬러를 각각 16진수로 표현 (R,G,B)
# font.color.rgb = RGBColor(0xFF,0x24,0xE9)



# # # for 문을 이용해서 
# # for run in para1:
# #     run.font.size = Pt(20)

# doc.save('text1.docx')

# doc = Document('예제 문서.docx')

# # 왼쪽정렬 
# paragraph1 = doc.paragraphs[1]
# paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT

# # 가운데 정렬
# paragraph2 = doc.paragraphs[2]
# paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# # 오른쪽 정렬
# paragraph3 = doc.paragraphs[3]
# paragraph3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# # 양쪽 정렬
# paragraph4 = doc.paragraphs[4]
# paragraph4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# # 텍스트 배분 (글자를 흩어서 배치)
# paragraph_last = doc.paragraphs[-1] # 마지막 문단
# paragraph_last.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
# paragraph_last.alignment = WD_PARAGRAPH_ALIGNMENT.DITRIBUTE # 컴파일 에러남

# 현재 작업경로에 저장 
# doc.save('정렬 예제.docx')

# 테이블 셀 정렬 
# # LEFT : 왼쪽정렬,  CENTER :  가운데 정렬, RIGHT : 오른쪽정렬 
# doc.tables[0].rows[0].cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT
# doc.tables[0].rows[0].cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
# doc.tables[0].rows[0].cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.RIGHT


# # doc.save('TryToc1.docx')

# 목차 생성 
# word = win32com.client.DispatchEx("Word.Application")
# word.Visible = False
# doc = word.Documents.Open('C:/PYTHONWORKSPACE/new_try.docx', Encoding='gbk')
# doc.TablesofContents(1).Update()
# doc.Close(SaveChange=True)
# word.Quit()


# # pdf 파일로 변환
# doc = aw.Document('정렬 예제.docx')

# doc.save('정렬예제.pdf')



# doc.save('정렬 예제.pdf')
