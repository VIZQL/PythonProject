# 가장 기본적인 기능(문서열기, 저장, 글자쓰기 등등)
from docx import Document

# 문단 정렬 
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 테이블 정렬 
from docx.enum.table import WD_TABLE_ALIGNMENT

# 테이블 셀 색 넣기
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# 문자 스타일 변경 
from docx.enum.style import WD_STYLE_TYPE

# 이미지 삽입
# Cm 과 Inch 단위를 사용하기 위한 모듈 
from docx.shared import Cm,Inches

# 한글 폰트 선 문단 입력, 후 폰트적용을 위한 import 
from docx.oxml.ns import qn

# 문자 크기를 변경하기 위해 폰트 크기와 관련된 class 를 import 
from docx.shared import Pt

# RGB color 를 사용하는 메쏘드
from docx.shared import RGBColor

# word 문서를 PDF 로 바꾸기 위한 모듈 
# import aspose.words as aw
# pip install aspose-words

# 목차를 만들기 위한 import 
import win32com.client
import os
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 현재 날짜 가져오기
# import datetime
# from datetime import datetime

# 내 모듈 import 하기
from Kangwon_News import*
from Kangwon_stocks_info import*
from stock_report import*
from Kangwon_schedule import*
from Kangwon_raw import*
# from kiwoom_pro import Kiwoom_condition # 패키지에서 모듈 임포트 
from Kiwoom_condition import*


import Kangwon_stocks_info
import stock_report
import Kangwon_News
import Kangwon_schedule
import Kangwon_raw
import Kiwoom_condition
# import kiwoom_pro/Kiwwom_condition 
# -*- coding: utf-8 -*



# 하이퍼 링크를 위한 함수 
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


# 셀 컬러 정의하기 

def cell_color(firstrow):
    for i in range (len(firstrow)):
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="D6E8F6"/>'.format(nsdecls('w')))
        firstrow[i]._tc.get_or_add_tcPr().append(shading_elm_1)



# 원하는 종목 특징주에서 뽑아오고 특징주 list 에서 삭제 

def stock_news_return(name, news_list, link_list):
    name_re = re.compile(name)
    news_result = []
    link_result = []

    for index, news in enumerate(news_list):
        m = name_re.search(news) # 주어진 문자열 중에 일치하는게 있는지 확인  
        
        if m:
            news_result.append(news_list[index])
            link_result.append(link_list[index])
            news_list.pop(index)
            link_list.pop(index)
        else:
            continue

    if(news_result): # 매칭 종목이 있을때 
        return news_result, link_result
    else: # 매칭 종목이 없을 때 
        return [""], [""]

# 새 워드 문서 만들기 
# doc = Document()

doc = Document('C:\PYTHONWORKSPACE\\TP2.docx')

today_time = datetime.date.today().strftime("%Y%m%d")  

# 헤딩 폰트 사이즈 변경
font = doc.styles['Title'].font
font.name = 'Times New Roman'
font.size = Pt(20)

font = doc.styles['Heading 1'].font
font.name = 'Times New Roman'
font.color.rgb = RGBColor(0,0,0)


# 글자 사이즈 변경 
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(8)

# 1. 제목 넣기 
doc.add_heading('kikawo SD report_{} '.format(today_time), level = 0)

# 2. 문단 추가하기 
p = doc.add_paragraph()
# p.add_run('kikawo SD 리포트는 저작권의 보호를 받습니다. 작성자 허가없이 다른 사람에게 공유는 엄격히 금지됩니다. 지정된 수신자 외에 다른 사람에게 전달되는 것이 적발될 경우 민형사상의 책임을 질 수 있습니다.').font.size = Pt(8)
p.add_run('kikawo SD 리포트는 저작권의 보호를 받습니다. 작성자 허가없이 다른 사람에게 공유는 엄격히 금지됩니다. 지정된 수신자 외에 다른 사람에게 전달되는 것이 적발될 경우 민형사상의 책임을 질 수 있습니다. 리포트 관련 문의사항은 kikawo@naver.com 으로 연락바랍니다').underline = True
p.runs[0].font.size = Pt(8)

paragraph = doc.add_paragraph()
run = paragraph.add_run()

fldChar = OxmlElement('w:fldChar')  # creates a new element
fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:t')
fldChar3.text = "Right-click to update field."
fldChar2.append(fldChar3)

fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')

r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar4)
p_element = paragraph._p

## 다음페이지로 넘어가기 
doc.add_page_break()

# 3. Daily comment 넣기 
doc.add_heading('1. Daily Comment : ', level = 1)


# 4. 주요 국가 지수
doc.add_heading('2. 주요 국가 지수 : ', level = 1)


# 5. 시장 주도 종목 
doc.add_heading('3. 시장 주도 종목 : ', level = 1)


# 6. 상한가 특징종목 정리 

doc.add_heading('4. 상한가 특징종목 정리 : ', level = 1)


# 12. 증권사 리포트
doc.add_heading('10. 증권사 리포트: ', level = 1)


###################################### 13. 환율/원자재 가격등##############################################
doc.add_heading('11. 환율/원자재 가격: ', level = 1)


############################################# 달러 인덱스

# 표 삽입 - 1행 5열의 표 만들기 
table = doc.add_table(rows = 1, cols = 3)

table.style = doc.styles['Table Grid']
table.autofit = False
table.allow_autofit = False


table.columns[0].width = Cm(4)
table.columns[1].width = Cm(2)
table.columns[2].width = Cm(2)

# 표의 첫 행을 리스트로 가져오기 
first_row = table.rows[0].cells

first_row[0].width = Cm(4)
first_row[1].width = Cm(2)
first_row[2].width = Cm(2)

first_row[0].text = '분류'
first_row[1].text = '가격'
first_row[2].text = '변동(%)'

cell_color(first_row)

try: 
    currency, d_index = Kangwon_stocks_info.scrape_major_indice_money()
        
    data = table.add_row().cells
    data[0].text = "doller index"
    data[1].text = d_index[0]

    data[2].text = d_index[1]
    if (float(d_index[1].strip("%")) > 0):
        data[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0x00,0x00)
    elif (float(d_index[1].strip("%")) < 0):
        data[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x00,0xFF)
    else:
        pass
except:
    pass

############################################# 환율 

doc.add_paragraph('\n') # 한칸 띄우고 

# 표 삽입 - 1행 5열의 표 만들기 
table = doc.add_table(rows = 1, cols = 6)

table.style = doc.styles['Table Grid']
table.autofit = False
table.allow_autofit = False


table.columns[0].width = Cm(2)
table.columns[1].width = Cm(2)
table.columns[2].width = Cm(2)
table.columns[3].width = Cm(2)
table.columns[4].width = Cm(2)
table.columns[5].width = Cm(2)


# 표의 첫 행을 리스트로 가져오기 
first_row = table.rows[0].cells

first_row[0].width = Cm(2)
first_row[1].width = Cm(2)
first_row[2].width = Cm(2)
first_row[3].width = Cm(2)
first_row[4].width = Cm(2)
first_row[5].width = Cm(2)


first_row[0].text = '분류'
first_row[1].text = '가격'
first_row[2].text = '일간'
first_row[3].text = '1주간'
first_row[4].text = '1달간'
first_row[5].text = 'YTD'


cell_color(first_row)

try: 
    currency, d_index = Kangwon_stocks_info.scrape_major_indice_money()
        
    data = table.add_row().cells
    data[0].text = "달러/원"
    data[1].text = currency[0] # 가격 

    data[2].text = currency[1]
    if (float(currency[1].strip("%")) > 0):
        data[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0x00,0x00)
    elif (float(currency[1].strip("%")) < 0):
        data[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x00,0xFF)
    else:
        pass

    data[3].text = currency[2]
    if (float(currency[2].strip("%")) > 0):
        data[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0x00,0x00)
    elif (float(currency[2].strip("%")) < 0):
        data[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x00,0xFF)
    else:
        pass

    data[4].text = currency[3]
    if (float(currency[3].strip("%")) > 0):
        data[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0x00,0x00)
    elif (float(currency[3].strip("%")) < 0):
        data[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x00,0xFF)
    else:
        pass

    data[5].text = currency[4]
    if (float(currency[4].strip("%")) > 0):
        data[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0x00,0x00)
    elif (float(currency[4].strip("%")) < 0):
        data[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x00,0xFF)
    else:
        pass

except:
    pass


############################################# 원자재 

doc.add_paragraph('\n') # 한칸 띄우고 

# 표 삽입 - 1행 5열의 표 만들기 
table = doc.add_table(rows = 1, cols = 5)

# 만든 표의 스타일을 가장 기본 스타일인 'Table Grid' 로 설정 

table.style = doc.styles['Table Grid']
table.autofit = False
table.allow_autofit = False

table.columns[0].width = Cm(1)
table.columns[1].width = Cm(1)
table.columns[2].width = Cm(1)
table.columns[3].width = Cm(1)
table.columns[4].width = Cm(1)

# 폰트 사이즈 재 변경 
style.font.size = Pt(10)

# 표의 첫 행을 리스트로 가져오기 
first_row = table.rows[0].cells

first_row[0].width = Cm(3)
first_row[1].width = Cm(2)
first_row[2].width = Cm(2)
first_row[3].width = Cm(2)
first_row[4].width = Cm(2)

first_row[0].text = '상품'
first_row[1].text = '일간'
first_row[2].text = '1주간'
first_row[3].text = '1개월간'
first_row[4].text = 'YTD'

# 셀에 배경색 넣기
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

cell_color(first_row)

# print(len(first_row))
try: 
    Kangwon_raw.scrape_rawM()
    for i in range(len(Kangwon_raw.name_list)):
        
        data = table.add_row().cells
        data[0].text = Kangwon_raw.name_list[i]
        data[1].text = Kangwon_raw.day_rate_list[i]
        if (float(Kangwon_raw.day_rate_list[i].strip("%")) > 0):
            data[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0x00,0x00)
        elif (float(Kangwon_raw.day_rate_list[i].strip("%")) < 0):
            data[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x00,0xFF)
        else:
            pass
        data[2].text = Kangwon_raw.week_rate_list[i]
        if (float(Kangwon_raw.week_rate_list[i].strip("%")) > 0):
            data[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0x00,0x00)
        elif (float(Kangwon_raw.week_rate_list[i].strip("%")) < 0):
            data[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x00,0xFF)
        else:
            pass
        data[3].text = Kangwon_raw.month_rate_list[i]
        if (float(Kangwon_raw.month_rate_list[i].strip("%")) > 0):
            data[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0x00,0x00)
        elif (float(Kangwon_raw.month_rate_list[i].strip("%")) < 0):
            data[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x00,0xFF)
        else:
            pass
        data[4].text = Kangwon_raw.YTD_rate_list[i]
        if (float(Kangwon_raw.YTD_rate_list[i].strip("%")) > 0):
            data[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0x00,0x00)
        elif (float(Kangwon_raw.YTD_rate_list[i].strip("%")) < 0):
            data[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x00,0xFF)
        else:
            pass
except:
    pass


#  현재 작업경로에 저장 
doc.save('C:\\PYTHONWORKSPACE\\kikawo_SD_리포트_{}_1.docx'.format(today_time))


word = win32com.client.DispatchEx("Word.Application")
doc = word.Documents.Open('C:\\PYTHONWORKSPACE\\kikawo_SD_리포트_{}_1.docx'.format(today_time))
doc.TablesOfContents(1).Update()
doc.Close(SaveChanges=True)
word.Quit()
